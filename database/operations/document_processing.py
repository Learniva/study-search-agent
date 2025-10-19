"""
Document Processing for L2 Vector Store

Handles document upload, text extraction, chunking, and indexing into pgvector.
Supports PDF, DOCX, TXT, and MD files with intelligent chunking.

Uses Docling for advanced document understanding:
- Superior layout analysis
- Table extraction and preservation
- Structure-aware parsing
- Better handling of complex PDFs
"""

import os
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

# Docling - Advanced document understanding (Primary)
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
    from docling.chunking import HybridChunker  # Advanced chunking for RAG
    from docling_core.types.doc import DoclingDocument
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    
# PyMuPDF - Fallback for simple PDFs
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# DOCX processing
try:
    import docx2txt
    DOCX_AVAILABLE = True
except ImportError:
    try:
        from docx import Document
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False

# LangChain for text splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Google GenAI SDK for 768D embeddings (direct API)
import google.generativeai as genai

# Database operations
from sqlalchemy.orm import Session
from database.operations.rag import store_document_vectors, delete_document_vectors

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process documents for L2 Vector Store indexing.
    
    Handles:
    - Advanced text extraction using Docling (tables, layout, structure)
    - Fallback to PyMuPDF for simple PDFs
    - Intelligent chunking with overlap
    - Embedding generation via Google Gemini
    - Storage in PostgreSQL + pgvector
    """
    
    def __init__(
        self,
        chunk_size: int = 800,  # Reduced from 1000 to create more chunks (40-70+)
        chunk_overlap: int = 150,  # Reduced proportionally
        embedding_model: str = "models/embedding-001",
        use_docling: bool = True
    ):
        """
        Initialize document processor with Google Gemini 768D embeddings.
        
        Configured for optimal RAG performance:
        - chunk_size=800: Creates 40-70+ chunks for typical documents
        - HybridChunker: Keeps tables intact, preserves context
        - Automatic indexing: Triggered on upload
        
        Args:
            chunk_size: Size of text chunks in characters (default: 800 for 40-70+ chunks)
            chunk_overlap: Overlap between chunks for context continuity
            embedding_model: Google Gemini embedding model (default: models/embedding-001, 768D)
            use_docling: Whether to use Docling for advanced PDF parsing (default: True)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_model = embedding_model
        self.use_docling = use_docling and DOCLING_AVAILABLE
        
        # Initialize Docling converter if available
        if self.use_docling:
            try:
                # Configure Docling for optimal performance
                pipeline_options = PdfPipelineOptions()
                pipeline_options.do_table_structure = True  # Enable table extraction
                pipeline_options.do_ocr = False  # Disable OCR for speed (enable if needed)
                
                self.docling_converter = DocumentConverter(
                    allowed_formats=[InputFormat.PDF],
                    pipeline_options=pipeline_options
                )
                
                # Initialize HybridChunker for intelligent document chunking
                # This keeps tables intact, preserves headings with text, and respects paragraph boundaries
                # Configured to create 40-70+ chunks for typical documents
                self.hybrid_chunker = HybridChunker(
                    tokenizer=None,  # Use default tokenizer
                    max_tokens=self.chunk_size // 4,  # ~200 tokens per chunk (800/4)
                )
                logger.info("✅ Docling initialized with HybridChunker (chunk_size=800 for 40-70+ chunks)")
            except Exception as e:
                logger.warning(f"⚠️ Docling initialization failed: {e}, falling back to PyMuPDF")
                self.use_docling = False
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Configure Google GenAI SDK for 768D embeddings
        try:
            api_key = os.getenv("GOOGLE_API_KEY")
            if not api_key:
                raise ValueError("GOOGLE_API_KEY not found in environment")
            
            genai.configure(api_key=api_key)
            
            # Verify model access with test embedding
            test_result = genai.embed_content(
                model=self.embedding_model,
                content="test",
                task_type="retrieval_document"
            )
            embedding_dim = len(test_result['embedding'])
            
            logger.info(f"✅ Google Gemini embeddings initialized: {embedding_model} ({embedding_dim}D)")
            
            if embedding_dim != 768:
                logger.warning(f"⚠️  Expected 768D, got {embedding_dim}D")
                
        except Exception as e:
            logger.error(f"❌ Failed to initialize Google Gemini embeddings: {e}")
            raise
    
    def extract_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from document based on file type.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (extracted text content, extraction metadata)
            
        Raises:
            ValueError: If file type is unsupported
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            extraction_metadata = {
                'file_type': extension.lstrip('.'),
                'file_name': file_path.name,
                'extraction_method': 'unknown',
                'has_tables': False,
                'table_count': 0
            }
            
            if extension == '.pdf':
                text, pdf_metadata = self._extract_pdf(file_path)
                extraction_metadata.update(pdf_metadata)
                return text, extraction_metadata
            elif extension == '.docx':
                text = self._extract_docx(file_path)
                extraction_metadata['extraction_method'] = 'docx2txt'
                return text, extraction_metadata
            elif extension in ['.txt', '.md']:
                text = self._extract_text(file_path)
                extraction_metadata['extraction_method'] = 'plain_text'
                return text, extraction_metadata
            else:
                raise ValueError(f"Unsupported file type: {extension}")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using Docling (primary) or PyMuPDF (fallback).
        
        Docling provides:
        - Better layout analysis
        - Table extraction and preservation
        - Structure-aware parsing
        - Handling of complex multi-column layouts
        
        Returns:
            Tuple of (extracted text, metadata about extraction)
        """
        # Try Docling first for advanced parsing
        if self.use_docling:
            try:
                return self._extract_pdf_with_docling(file_path)
            except Exception as e:
                logger.warning(f"⚠️ Docling extraction failed: {e}, falling back to PyMuPDF")
        
        # Fallback to PyMuPDF for simple text extraction
        return self._extract_pdf_with_pymupdf(file_path)
    
    def _extract_pdf_with_docling(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using Docling for advanced document understanding.
        
        Features:
        - Layout-aware text extraction
        - Table detection and formatting  
        - Structure preservation (headings, lists, etc.)
        - Multi-column handling
        - Returns DoclingDocument for HybridChunker
        
        Returns:
            Tuple of (extracted text, metadata about extraction)
        """
        try:
            logger.info(f"📄 Using Docling with HybridChunker for advanced PDF parsing: {file_path.name}")
            
            # Convert document using Docling
            result = self.docling_converter.convert(str(file_path))
            
            # Get the DoclingDocument object (needed for HybridChunker)
            doc: DoclingDocument = result.document
            
            # Count tables
            table_count = 0
            if hasattr(doc, 'tables') and doc.tables:
                table_count = len(doc.tables)
                logger.info(f"📊 Detected {table_count} tables in PDF")
            
            # Export to markdown for storage
            # HybridChunker will handle this document intelligently later
            markdown_text = doc.export_to_markdown()
            
            if not markdown_text or not markdown_text.strip():
                raise ValueError("Docling extracted empty content")
            
            # Store the DoclingDocument for chunking later
            # We'll use it with HybridChunker instead of RecursiveCharacterTextSplitter
            self._current_docling_doc = doc  # Cache for chunking
            
            # Build metadata
            metadata = {
                'extraction_method': 'docling_hybrid',
                'has_tables': table_count > 0,
                'table_count': table_count,
                'char_count': len(markdown_text),
                'structure_preserved': True,
                'markdown_format': True,
                'uses_hybrid_chunker': True
            }
            
            logger.info(f"✅ Docling extraction successful: {len(markdown_text)} chars, {table_count} tables")
            return markdown_text, metadata
            
        except Exception as e:
            logger.error(f"Docling PDF extraction failed: {e}")
            raise
    
    def _extract_pdf_with_pymupdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using PyMuPDF (fallback method).
        
        Returns:
            Tuple of (extracted text, metadata about extraction)
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available. Install with: pip install pymupdf")
        
        try:
            logger.info(f"📄 Using PyMuPDF for PDF extraction: {file_path.name}")
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num}]\n{text}")
            
            doc.close()
            result = "\n\n".join(text_parts)
            
            # Build metadata
            metadata = {
                'extraction_method': 'pymupdf',
                'has_tables': False,  # PyMuPDF doesn't do table detection
                'table_count': 0,
                'char_count': len(result),
                'structure_preserved': False,
                'markdown_format': False
            }
            
            logger.info(f"✅ PyMuPDF extraction successful: {len(result)} chars")
            return result, metadata
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}")
            raise
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing not available. Install with: pip install python-docx or docx2txt")
        
        try:
            # Try docx2txt first (more reliable)
            text = docx2txt.process(str(file_path))
            if text and text.strip():
                return text
        except:
            pass
        
        try:
            # Fallback to python-docx
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from TXT or MD file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def chunk_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks with metadata.
        
        Uses HybridChunker for Docling documents (keeps tables intact, preserves structure).
        Falls back to RecursiveCharacterTextSplitter for other formats.
        
        Args:
            text: Full text to chunk
            metadata: Optional metadata to attach to each chunk
            
        Returns:
            List of chunks with content and metadata
        """
        # Check if we should use HybridChunker (for Docling documents)
        use_hybrid = (
            self.use_docling and 
            metadata and 
            metadata.get('extraction_method') == 'docling_hybrid' and
            hasattr(self, '_current_docling_doc')
        )
        
        if use_hybrid:
            # Use HybridChunker for intelligent document-aware chunking
            try:
                logger.info("📝 Using HybridChunker for intelligent document chunking")
                doc = self._current_docling_doc
                
                # Chunk the DoclingDocument
                chunk_iter = self.hybrid_chunker.chunk(doc)
                chunks_list = list(chunk_iter)
                
                result = []
                for idx, chunk in enumerate(chunks_list):
                    # Extract text from the chunk
                    chunk_text = chunk.text if hasattr(chunk, 'text') else str(chunk)
                    
                    chunk_dict = {
                        'content': chunk_text,
                        'chunk_index': idx,
                        'metadata': {
                            **(metadata or {}),
                            'chunk_number': idx + 1,
                            'total_chunks': len(chunks_list),
                            'char_count': len(chunk_text),
                            'chunking_method': 'hybrid_chunker'
                        }
                    }
                    result.append(chunk_dict)
                
                # Clean up cached document
                delattr(self, '_current_docling_doc')
                
                chunk_count = len(result)
                logger.info(f"✅ HybridChunker created {chunk_count} context-aware chunks")
                
                # Provide feedback on chunk count
                if chunk_count < 40:
                    logger.info(f"   ℹ️  Document is relatively short ({chunk_count} chunks)")
                elif 40 <= chunk_count <= 70:
                    logger.info(f"   ✅ Optimal chunk count for RAG ({chunk_count} chunks)")
                elif chunk_count > 70:
                    logger.info(f"   ℹ️  Large document with {chunk_count} chunks")
                
                return result
                
            except Exception as e:
                logger.warning(f"⚠️ HybridChunker failed: {e}, falling back to RecursiveCharacterTextSplitter")
                # Fall through to standard chunking
                if hasattr(self, '_current_docling_doc'):
                    delattr(self, '_current_docling_doc')
        
        # Standard chunking for non-Docling documents or fallback
        chunks = self.text_splitter.split_text(text)
        
        result = []
        for idx, chunk_text in enumerate(chunks):
            chunk = {
                'content': chunk_text,
                'chunk_index': idx,
                'metadata': {
                    **(metadata or {}),
                    'chunk_number': idx + 1,
                    'total_chunks': len(chunks),
                    'char_count': len(chunk_text),
                    'chunking_method': 'recursive_text_splitter'
                }
            }
            result.append(chunk)
        
        return result
    
    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        Generate 768D embeddings using Google Gemini.
        
        Args:
            texts: List of text strings to embed
            
        Returns:
            List of 768-dimensional embedding vectors
        """
        try:
            # Google GenAI SDK doesn't have async batch embed, so use sync
            embeddings = []
            for text in texts:
                result = genai.embed_content(
                    model=self.embedding_model,
                    content=text,
                    task_type="retrieval_document"
                )
                embeddings.append(result['embedding'])
            
            logger.debug(f"Generated {len(embeddings)} embeddings ({len(embeddings[0])}D each)")
            return embeddings
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise
    
    def generate_embeddings_sync(self, texts: List[str]) -> List[List[float]]:
        """
        Generate 768D embeddings synchronously (for background tasks).
        
        Args:
            texts: List of text strings to embed
            
        Returns:
            List of 768-dimensional embedding vectors
        """
        try:
            embeddings = []
            for text in texts:
                result = genai.embed_content(
                    model=self.embedding_model,
                    content=text,
                    task_type="retrieval_document"
                )
                embeddings.append(result['embedding'])
            
            logger.debug(f"Generated {len(embeddings)} embeddings ({len(embeddings[0])}D each)")
            return embeddings
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise
    
    async def process_and_index_document(
        self,
        db: Session,
        file_path: str,
        document_name: str,
        user_id: Optional[str] = None,
        course_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Complete pipeline: extract, chunk, embed, and index document.
        
        Uses Docling for advanced PDF parsing with table extraction.
        
        Args:
            db: Database session
            file_path: Path to document file
            document_name: Name of the document
            user_id: Optional user ID for ownership
            course_id: Optional course ID
            
        Returns:
            Dictionary with indexing results including extraction metadata
        """
        try:
            # Generate unique document ID
            document_id = str(uuid.uuid4())
            
            # Extract text with Docling (returns text + metadata)
            logger.info(f"Extracting text from {document_name}...")
            text, extraction_metadata = self.extract_text(file_path)
            
            if not text or len(text.strip()) < 10:
                raise ValueError("Document is empty or too short")
            
            # Chunk text with extraction metadata
            file_type = Path(file_path).suffix.lstrip('.')
            metadata = {
                'document_name': document_name,
                'document_type': file_type,
                'file_size': os.path.getsize(file_path),
                **extraction_metadata  # Include Docling extraction info
            }
            
            logger.info(f"Chunking text ({len(text)} chars, method: {extraction_metadata.get('extraction_method')})...")
            logger.info(f"Target: 40-70+ chunks (chunk_size={self.chunk_size}, overlap={self.chunk_overlap})")
            chunks = self.chunk_text(text, metadata)
            
            # Generate embeddings
            logger.info(f"Generating embeddings for {len(chunks)} chunks...")
            chunk_texts = [chunk['content'] for chunk in chunks]
            embeddings = await self.generate_embeddings(chunk_texts)
            
            # Attach embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                chunk['embedding'] = embedding
                chunk['type'] = file_type
            
            # Store in database
            logger.info(f"Storing {len(chunks)} vectors in database...")
            vectors_stored = store_document_vectors(
                db=db,
                document_id=document_id,
                document_name=document_name,
                chunks=chunks,
                user_id=user_id,
                course_id=course_id
            )
            
            # Log indexing results with chunk count analysis
            chunk_count = len(chunks)
            logger.info(f"✅ Document indexed: {document_name}")
            logger.info(f"   📊 Chunks: {chunk_count} | Vectors: {vectors_stored}")
            logger.info(f"   📝 Method: {extraction_metadata.get('extraction_method')}")
            logger.info(f"   🔧 Chunking: {extraction_metadata.get('chunking_method', 'unknown')}")
            
            # Warn if chunk count is outside optimal range
            if chunk_count < 40:
                logger.warning(f"⚠️  Low chunk count ({chunk_count}). Expected 40-70+ for typical documents.")
                logger.warning(f"   This might reduce RAG accuracy. Document may be too short.")
            elif chunk_count > 200:
                logger.warning(f"⚠️  High chunk count ({chunk_count}). Consider increasing chunk_size.")
            else:
                logger.info(f"✅ Chunk count in optimal range: {chunk_count}")
            
            return {
                'success': True,
                'document_id': document_id,
                'document_name': document_name,
                'chunks_created': chunk_count,
                'vectors_stored': vectors_stored,
                'total_chars': len(text),
                'extraction_info': {
                    **extraction_metadata,
                    'chunk_size_config': self.chunk_size,
                    'chunk_overlap_config': self.chunk_overlap
                }
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'document_name': document_name
            }
    
    def process_and_index_document_sync(
        self,
        db: Session,
        file_path: str,
        document_name: str,
        user_id: Optional[str] = None,
        course_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Synchronous version of process_and_index_document (for background tasks).
        
        Uses Docling for advanced PDF parsing with table extraction.
        
        Args:
            db: Database session
            file_path: Path to document file
            document_name: Name of the document
            user_id: Optional user ID for ownership
            course_id: Optional course ID
            
        Returns:
            Dictionary with indexing results including extraction metadata
        """
        try:
            # Generate unique document ID
            document_id = str(uuid.uuid4())
            
            # Extract text with Docling (returns text + metadata)
            logger.info(f"Extracting text from {document_name}...")
            text, extraction_metadata = self.extract_text(file_path)
            
            if not text or len(text.strip()) < 10:
                raise ValueError("Document is empty or too short")
            
            # Chunk text with extraction metadata
            file_type = Path(file_path).suffix.lstrip('.')
            metadata = {
                'document_name': document_name,
                'document_type': file_type,
                'file_size': os.path.getsize(file_path),
                **extraction_metadata  # Include Docling extraction info
            }
            
            logger.info(f"Chunking text ({len(text)} chars, method: {extraction_metadata.get('extraction_method')})...")
            logger.info(f"Target: 40-70+ chunks (chunk_size={self.chunk_size}, overlap={self.chunk_overlap})")
            chunks = self.chunk_text(text, metadata)
            
            # Generate embeddings (sync)
            logger.info(f"Generating embeddings for {len(chunks)} chunks...")
            chunk_texts = [chunk['content'] for chunk in chunks]
            embeddings = self.generate_embeddings_sync(chunk_texts)
            
            # Attach embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                chunk['embedding'] = embedding
                chunk['type'] = file_type
            
            # Store in database
            logger.info(f"Storing {len(chunks)} vectors in database...")
            vectors_stored = store_document_vectors(
                db=db,
                document_id=document_id,
                document_name=document_name,
                chunks=chunks,
                user_id=user_id,
                course_id=course_id
            )
            
            # Log indexing results with chunk count analysis
            chunk_count = len(chunks)
            logger.info(f"✅ Document indexed: {document_name}")
            logger.info(f"   📊 Chunks: {chunk_count} | Vectors: {vectors_stored}")
            logger.info(f"   📝 Method: {extraction_metadata.get('extraction_method')}")
            logger.info(f"   🔧 Chunking: {extraction_metadata.get('chunking_method', 'unknown')}")
            
            # Warn if chunk count is outside optimal range
            if chunk_count < 40:
                logger.warning(f"⚠️  Low chunk count ({chunk_count}). Expected 40-70+ for typical documents.")
                logger.warning(f"   This might reduce RAG accuracy. Document may be too short.")
            elif chunk_count > 200:
                logger.warning(f"⚠️  High chunk count ({chunk_count}). Consider increasing chunk_size.")
            else:
                logger.info(f"✅ Chunk count in optimal range: {chunk_count}")
            
            return {
                'success': True,
                'document_id': document_id,
                'document_name': document_name,
                'chunks_created': chunk_count,
                'vectors_stored': vectors_stored,
                'total_chars': len(text),
                'extraction_info': {
                    **extraction_metadata,
                    'chunk_size_config': self.chunk_size,
                    'chunk_overlap_config': self.chunk_overlap
                }
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'document_name': document_name
            }


# Global processor instance
_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get or create global document processor instance."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor


def remove_document_from_vector_store(
    db: Session,
    document_name: str
) -> bool:
    """
    Remove all vectors for a document from the vector store.
    
    Args:
        db: Database session
        document_name: Name of document to remove
        
    Returns:
        True if successful
    """
    try:
        # Find all vectors with this document name
        from database.models import DocumentVector
        vectors = db.query(DocumentVector).filter(
            DocumentVector.document_name == document_name
        ).all()
        
        if not vectors:
            logger.warning(f"No vectors found for document: {document_name}")
            return True
        
        # Delete all vectors
        for vector in vectors:
            db.delete(vector)
        
        db.commit()
        logger.info(f"✅ Removed {len(vectors)} vectors for: {document_name}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to remove document vectors: {e}")
        db.rollback()
        return False


__all__ = [
    'DocumentProcessor',
    'get_document_processor',
    'remove_document_from_vector_store'
]

